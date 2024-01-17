# This program will first extract the text from PDF file then translate the text. The final output is in DOCX format
# The original text will be presented alongside with translated text for comparison purposes

import PyPDF2
import time
from docx import Document
from googletrans import Translator
import os

## FUNCTION FOR EXTRACTION AND TRANSLATING TEXT ##
def textTRANSLATION():
    translator = Translator()
    document = Document()
    document.add_heading(PDFFILE, 0)

    pdfFileObj = open(PDFFILE, 'rb')
    pdfReader = PyPDF2.PdfReader(pdfFileObj)

    ## ESTABLISHING NUMBER OF PAGES WITHIN PDF FILE ##
    numPages = len(pdfReader.pages)
    # print(numPages)

    ## Translates text and outputs to DOCX. Currently auto detects language and defaults to english as output ##
    for i in range(numPages):
        try:
            pageObj = pdfReader.pages[i]
            TEXT = pageObj.extract_text()
            print(f'Page: {i + 1} translation and saving to DOCX in progress')
            translation = translator.translate(TEXT)
            Input = translation.text
            q = document.add_paragraph()
            heading = q.add_run(f'Page: {i + 1}')
            heading.bold = True
            r = document.add_paragraph(TEXT)
            p = document.add_paragraph(Input)
            document.save(f'{PDFFILE}.docx')
            time.sleep(1)
        except:
            print(f'Error occurred on page{i + 1}')

## SET LOCATION OF PDFs TO BE TRANSLATED (must be same as program file location)
folder_path = r'C:\Python\Woolworths Specials'
list_of_files = os.listdir(folder_path)

## LOOPING THROUGH ALL PDFS IN FOLDER LOCATION AND TRANSLATING
for file_name in list_of_files:
    if '.pdf' in file_name:
        PDFFILE = file_name
        print("Translation in progress .......", file_name)
        textTRANSLATION()

print("TEXT EXTRACT and Translation COMPLETE")